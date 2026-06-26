from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SRC = ROOT / "第三部分_基于GNN的压裂工况标签识别阶段总结_扩展版.docx"
OUT = ROOT / "第三部分_基于GNN的压裂工况标签识别阶段总结_扩展加强版.docx"
TMP = ROOT / "_tmp_gnn_report_rich.docx"


def set_run_font(run, font_name: str = "宋体", size: float = 10.5, bold: bool = False) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = "Heading 1"
    run = p.add_run(text)
    set_run_font(run, "黑体", 14, True)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_run_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run)


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    TMP.write_bytes(SRC.read_bytes())
    doc = Document(str(TMP))
    doc.add_page_break()

    add_heading(doc, "3.13 阶段小结")
    add_para(
        doc,
        "总体来看，本阶段工作已经从单一模型训练推进到较完整的压裂工况智能识别技术链路。我们首先明确了字段口径，即以WORKING_TYPE作为唯一输出标签，以砂比、泵压、排量、粘度、累计液量、累计砂量、暂堵剂相关参数及其动态变化作为输入特征，排除了井号、段号、时间、备注、人工标记和结果性字段，保证模型学习的是施工参数与工况之间的关系，而不是表格标识信息。",
    )
    add_para(
        doc,
        "在数据组织方面，本阶段坚持以压裂段为基本建模单元。所有训练、验证、测试和迁移学习划分均按段进行，同一段的样本不会同时进入不同子集，从而避免时序泄漏。窗口构样采用1至6点历史窗口，实际结果表明，约40秒历史上下文对应的窗口4在多分类任务中表现最好，说明工况识别既需要短时趋势信息，也需要控制过长窗口带来的无关扰动。",
    )
    add_para(
        doc,
        "从模型迭代过程看，最初的多分类GNN暴露出典型的不平衡学习问题：正常样本占比过高，模型容易倾向输出正常。后续通过正常/异常二分类、异常样本权重、正常样本下采样、加砂后截断、动态特征工程和LightGBM强基线对照，逐步明确了问题根源。结果显示，模型效果提升主要来自数据口径和特征表达的改进，而不是单纯增加训练轮数或堆叠模型复杂度。",
    )
    add_para(
        doc,
        "新增便签数据对本阶段结果具有关键作用。原始数据中异常占比约1.25%，新增数据后合并数据异常占比提升到约5.00%，同时补充了缝内暂堵、缝高延伸、延伸受阻、滤失过大等原始数据中缺失或稀缺的标签。合并数据后的多分类窗口4测试Macro F1达到0.5739，虽然距离工程化自动判别仍有差距，但已经证明模型能够识别缝口暂堵、缝内暂堵等部分异常类型，不再只是简单学习正常类。",
    )
    add_para(
        doc,
        "需要强调的是，当前指标应分任务口径解释。当前工况多分类识别的测试Accuracy约80.70%，Macro F1约57.39%，反映的是直接识别当前WORKING_TYPE的能力；正常/异常二分类Accuracy优先版本测试Accuracy约83.89%，Macro F1约72.69%，说明二分类预警能力明显好于细分类识别；严格下一工况转移概率模型测试Accuracy达到99.47%，迁移集Accuracy达到97.03%，但该模型输入包含当前工况，业务含义是预测下一采样点状态转移概率，适合作为短时风险趋势预警，而不能简单等同于当前点多分类识别。",
    )
    add_para(
        doc,
        "因此，本阶段可以形成一个比较清晰的汇报结论：在现有数据条件下，模型路线是可行的，GNN和动态特征能够支撑段内时序建模，LightGBM基线验证了特征工程的有效性，转移概率模型能够提供90%以上的短时状态预测准确率；但当前工况细分类识别仍受异常样本不足、标签长尾、段间分布差异和标签边界不一致限制，尚不适合作为无人值守的自动决策系统。",
    )
    add_para(
        doc,
        "后续工作不应只围绕“把指标调高”展开，而应建立数据、标签、模型和现场验证的闭环。特别是对于甲方汇报，需要明确说明：若要继续提升多分类识别能力，关键资源不是更多训练轮数，而是更多覆盖不同井段、不同施工阶段、不同异常类型的高质量标注数据，以及专家对典型误判样本和异常起止边界的复核。只有解决数据覆盖和标签一致性问题，模型精度才有稳定提升空间。",
    )
    add_bullets(
        doc,
        [
            "阶段成果：完成了数据接入、字段清洗、分段划分、窗口构样、GNN建模、LightGBM强基线、多分类恢复、转移概率预测和迁移学习验证。",
            "有效改进：新增异常数据、加砂后截断和动态特征工程是当前最有效的提升手段。",
            "当前能力：模型能够识别部分异常工况，二分类预警效果优于多分类细分，下一工况转移预测可达到90%以上准确率。",
            "主要短板：长尾异常样本不足，部分类别只出现在个别段，标签边界仍需专家统一口径。",
            "应用定位：现阶段更适合作为辅助标注、风险提示和专家复核工具，不建议直接作为自动控制决策依据。",
            "下一步重点：补充异常样本、复核标签边界、固定独立验证集、优化GNN结构，并将当前识别与转移概率预警联合输出。",
        ]
    )
    add_para(
        doc,
        "综上，本阶段工作已经验证了基于GNN的压裂工况标签识别路线具备继续推进价值。下一阶段建议以“异常样本补充和专家复核”为前置条件，以“当前工况识别+下一工况概率预警”为主要技术形态，以“独立井段验证指标”为验收依据，逐步推动模型从离线实验走向现场辅助应用。",
    )

    doc.save(str(TMP))
    OUT.write_bytes(TMP.read_bytes())
    TMP.unlink(missing_ok=True)
    print(OUT)


if __name__ == "__main__":
    main()
