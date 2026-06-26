from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
FIG_DIR = ROOT / "report_figures"
OUT_DOCX = ROOT / "第三部分_基于GNN的压裂工况标签识别阶段总结.docx"
OUT_MD = ROOT / "第三部分_基于GNN的压裂工况标签识别阶段总结.md"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_default_font(doc: Document) -> None:
    styles = doc.styles
    for style_name in ["Normal", "Body Text"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "宋体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            style = styles[style_name]
            style.font.name = "黑体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            style.font.color.rgb = RGBColor(0, 0, 0)


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("3  基于GNN的压裂工况标签识别研究")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}" if level <= 3 else "Normal"
    run = p.add_run(text)
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(14 if level == 1 else 12)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_picture(doc: Document, filename: str, caption: str, width: float = 5.8) -> None:
    path = FIG_DIR / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        cap_run.font.name = "宋体"
        cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        cap_run.font.size = Pt(9)


def build_docx() -> None:
    doc = Document()
    set_default_font(doc)
    section = doc.sections[0]
    section.page_width = 7560310
    section.page_height = 10692130
    section.left_margin = 720090
    section.right_margin = 720090
    section.top_margin = 720090
    section.bottom_margin = 720090

    add_title(doc)
    add_heading(doc, "3.1 研究目标与总体思路")
    add_para(doc, "本部分围绕压裂施工秒点数据的工况标签识别开展研究，目标是在已有专家规则标注和现场标签基础上，建立能够利用时间连续性、段内结构关系和动态工艺参数变化的智能识别模型。模型输出对象为 WORKING_TYPE，即压裂施工过程中的工况标签；输入对象为砂比、泵压、排量、粘度、累计液量等工艺数值特征及其时间变化特征，不将井名、段名、时间、备注、标记字段和结果性字段作为模型输入。")
    add_para(doc, "前期探索首先按原始多分类标签进行GNN建模，随后针对样本极度不均衡、异常标签长尾、部分异常只在个别段出现等问题，逐步增加正常/异常二分类、异常样本重权重、加砂后截断、动态特征、LightGBM强基线、多分类恢复、下一工况转移概率预测以及迁移学习实验。整体工作形成了从标签识别到风险概率输出的完整实验链路。")
    add_table(doc, ["模块", "核心问题", "当前实现"], [
        ["标签识别", "当前时刻属于哪类工况", "GNN与LightGBM多分类/二分类模型"],
        ["异常识别", "正常样本占比过高导致异常召回不足", "加砂后截断、类别权重、正常下采样"],
        ["转移预测", "下一时刻从当前工况转向哪类工况", "严格使用当前工况和历史窗口预测下一工况概率"],
        ["迁移学习", "新段或新井标签分布与训练集不一致", "源域预训练 + 迁移集少量样本微调"],
    ])
    doc.add_page_break()

    add_heading(doc, "3.2 数据基础与标签体系")
    add_para(doc, "数据来源包括原始分段压裂数据和后续新增的便签数据。原始数据共约16.40万条秒点记录，正常样本约16.19万条，异常样本约2055条，异常占比约1.25%。新增便签数据共3.38万条记录，异常样本约7828条，异常占比约23.16%，显著补充了原始数据中缺失或稀缺的异常类型。合并后总样本约19.78万条，异常样本约9883条，异常占比提升至约5.00%。")
    add_para(doc, "从标签体系看，正常工况仍占绝对多数；异常标签包含主缝延伸、缝口暂堵、缝内暂堵、砂堵、缝高延伸、其他、延伸受阻、滤失过大等。新增数据带来了缝内暂堵、缝高延伸、延伸受阻、滤失过大等工况，使多分类建模从原先的少数异常类型扩展为更完整的现场工况体系。")
    add_table(doc, ["数据范围", "总样本量", "正常样本", "异常样本", "异常占比"], [
        ["原始分段数据", "163955", "161900", "2055", "1.25%"],
        ["新增便签数据", "33800", "25972", "7828", "23.16%"],
        ["合并数据", "197755", "187872", "9883", "5.00%"],
    ])
    add_picture(doc, "01_normal_abnormal_ratio.png", "图3-1 正常与异常样本占比对比", 5.5)
    add_picture(doc, "02_abnormal_label_distribution.png", "图3-2 异常标签数量分布", 5.5)
    doc.add_page_break()

    add_heading(doc, "3.3 样本组织与GNN建模方式")
    add_para(doc, "压裂数据具有明显的段内连续性。每个压裂段是一条按时间排序的10秒采样序列，段与段之间不应直接拼接。因此，训练、验证、测试和迁移学习集均按段划分，同一段产生的所有样本只能进入一个子集，避免同一段的相邻时间点同时出现在训练和测试中造成时序泄漏。")
    add_para(doc, "GNN建模时可将一个压裂段表示为时间图结构：图中节点对应10秒采样点，节点特征为该时刻的工艺数值特征，时间相邻节点之间建立边。1至6点输入窗口对应不同历史上下文长度，本质上是在图结构或序列窗口中利用前序节点信息判断目标时刻工况。")
    add_table(doc, ["窗口长度", "输入内容", "预测目标", "业务含义"], [
        ["1点", "当前采样点特征", "当前WORKING_TYPE", "只看瞬时参数"],
        ["2-3点", "最近20-30秒特征", "目标点WORKING_TYPE", "引入短时变化趋势"],
        ["4点", "最近40秒特征", "目标点WORKING_TYPE", "当前最优窗口，兼顾趋势与噪声"],
        ["5-6点", "最近50-60秒特征", "目标点WORKING_TYPE", "更长上下文，可能引入无关波动"],
    ])
    add_para(doc, "在当前实验中，窗口长度4表现最好，说明约40秒的历史工况变化对标签识别更有效；窗口过短不足以体现动态变化，窗口过长则可能引入与当前工况无关的早期扰动。")
    doc.add_page_break()

    add_heading(doc, "3.4 类别不平衡与小样本问题诊断")
    add_para(doc, "模型初始效果不理想的核心原因不是模型结构单一，而是数据分布存在强烈不均衡和长尾问题。原始全量点级分布中异常占比仅约1.25%，模型即使大量预测正常，也可以获得较高表观准确率，但对异常工况的召回能力不足。因此，项目评估不能只看Accuracy，必须重点看Macro F1、异常类Precision、异常类Recall和混淆矩阵。")
    add_para(doc, "进一步统计发现，部分异常类型只出现在少数段甚至个别段中，这会导致按段划分后某些异常类只出现在训练集或只出现在测试/迁移集。此时模型面对的是典型的小样本、长尾、跨段分布漂移问题，单纯增加训练轮数或提高模型复杂度难以根本解决。")
    add_bullets(doc, [
        "正常样本过多会使损失函数被正常类主导，模型趋向输出正常。",
        "异常样本少且分布集中，会造成异常类边界学习不稳定。",
        "某些标签在训练集中样本不足，会导致测试或迁移场景中无法可靠识别。",
        "按段划分虽然降低泄漏风险，但会暴露真实部署中的跨段泛化困难。",
    ])
    add_picture(doc, "03_new_segment_label_heatmap.png", "图3-3 新增数据各段标签分布热力图", 5.8)
    doc.add_page_break()

    add_heading(doc, "3.5 数据处理与特征工程改进")
    add_para(doc, "针对加砂前大量正常样本对异常识别帮助有限的问题，本阶段实现了“加砂后截断”策略，即以砂比SB首次大于0作为有效施工识别起点，去除大量加砂前纯正常背景点。该策略不会改变原始异常标签本身，但会改变窗口样本统计，因为不同窗口长度、段级划分和截断位置都会影响可生成的训练样本数量。")
    add_para(doc, "同时加入动态特征，除原始工艺参数外，为泵压、排量、砂比、累计砂量、累计液量、粘度等参数构造一阶差分、二阶差分、滚动均值、滚动标准差、窗口内最大最小值、斜率和砂比正值比例等特征。该类特征能够表达参数的变化方向和变化速度，比单点静态值更贴近现场专家判断逻辑。")
    add_table(doc, ["处理方式", "目的", "实际效果"], [
        ["加砂后截断", "减少无效正常背景样本", "异常占比提升，二分类F1提高"],
        ["正常下采样", "降低正常类对损失函数的压制", "提升异常召回，但需控制误报"],
        ["类别权重", "少样本类别获得更高训练权重", "有一定帮助，但不是根本解决方案"],
        ["动态特征", "表达压力、排量、砂比等变化趋势", "对二分类和多分类均有贡献"],
    ])
    doc.add_page_break()

    add_heading(doc, "3.6 模型实验结果与阶段进展")
    add_para(doc, "模型实验从原始多分类GNN开始，逐步演进为多路径对比。原始GNN在全量不平衡数据上容易学成“正常类优先”的分类器；二分类GNN和LightGBM基线能够一定程度改善异常识别；加入加砂后截断与动态特征后，二分类最优测试Macro F1提升至约0.5533；进一步引入新增便签数据并恢复多分类后，最优窗口4的测试Macro F1达到0.5739。")
    add_para(doc, "从结果看，当前最有效的改进并非单纯更换模型，而是补充异常数据、改进样本口径和引入动态特征。特别是新增数据使异常类型更完整，显著缓解了原始数据中部分标签完全缺失的问题。")
    add_table(doc, ["实验路径", "任务形式", "代表结果", "结论"], [
        ["原始GNN", "多分类", "异常识别弱", "受正常类主导明显"],
        ["GNN二分类", "正常/异常", "Macro F1约0.517", "可作为异常预警初版"],
        ["LightGBM二分类", "正常/异常", "Macro F1约0.524", "强基线略优于早期GNN"],
        ["加砂后+动态特征", "二分类", "Macro F1约0.5533", "特征工程有效"],
        ["合并新增数据", "多分类", "Macro F1约0.5739", "数据补充带来最直接提升"],
    ])
    add_picture(doc, "04_model_macro_f1_comparison.png", "图3-4 不同实验路径Macro F1对比", 5.8)
    doc.add_page_break()

    add_heading(doc, "3.7 多分类识别效果分析")
    add_para(doc, "在合并数据多分类实验中，窗口4取得当前最优测试结果：测试集Accuracy为0.8070，Macro F1为0.5739。正常类F1达到0.8942，说明模型对主类识别稳定；缝口暂堵F1为0.6097，缝内暂堵F1为0.4978，说明新增数据中的暂堵类已经具备一定可学习性；主缝延伸Precision较高但Recall较低，说明模型预测该类时较谨慎，漏检较多。")
    add_para(doc, "需要注意，Macro F1按类别平均，不会被正常类样本数量淹没，因此更能反映异常标签识别的真实能力。当前结果说明模型已经从“几乎全预测正常”改善为能够识别部分异常类别，但长尾类别如其他、延伸受阻、滤失过大、缝高延伸在当前测试划分中样本不足或缺失，仍不能证明已具备稳定识别能力。")
    add_table(doc, ["类别", "Precision", "Recall", "F1", "测试样本数"], [
        ["正常", "0.8321", "0.9663", "0.8942", "4570"],
        ["主缝延伸", "0.9365", "0.1744", "0.2940", "1015"],
        ["缝内暂堵", "0.4130", "0.6264", "0.4978", "182"],
        ["缝口暂堵", "0.6157", "0.6038", "0.6097", "260"],
    ])
    add_picture(doc, "05_multiclass_per_class_metrics.png", "图3-5 多分类测试集各类别识别指标", 5.8)
    doc.add_page_break()

    add_heading(doc, "3.8 下一工况转移概率预测")
    add_para(doc, "为满足现场对风险提前量和工况演化趋势的需求，本阶段进一步建立严格的“从当前工况到下一工况”的预测模型。该模型输入为当前时刻及历史窗口内的工艺参数动态特征，同时显式输入当前工况标签y_t，输出下一采样点工况y_{t+1}的概率分布。该任务不是单纯识别当前标签，而是估计从一种工况转移到另一种工况的风险概率。")
    add_para(doc, "在该建模方式下，测试集Accuracy达到0.9947，Macro F1达到0.9846，迁移集Macro F1达到0.8106。该指标较高的原因在于压裂工况具有强连续性，绝大多数相邻10秒点保持同一状态；同时模型已经知道当前工况，因此更适合解释为短时转移风险模型，而不能直接与纯工况识别模型等同对比。")
    add_table(doc, ["当前工况", "下一工况概率示例", "业务解释"], [
        ["正常", "下一点仍为正常的平均概率约99.99%", "正常状态具有强连续性"],
        ["正常", "转向缝内暂堵、缝口暂堵存在低概率尾部风险", "可用于提前预警排序"],
        ["异常", "下一点维持同类异常概率通常较高", "反映异常持续性"],
        ["低频异常", "概率估计仍受样本不足限制", "需要更多案例校准"],
    ])
    add_picture(doc, "06_normal_to_abnormal_transition_probs.png", "图3-6 正常状态转向异常状态的预测概率", 5.8)
    doc.add_page_break()

    add_heading(doc, "3.9 迁移学习与跨段泛化")
    add_para(doc, "由于不同段、不同井和不同施工阶段的参数分布存在差异，基础模型在迁移集上的表现通常低于训练段和测试段。为验证少量目标域样本对模型适配的作用，本阶段实现了源域预训练加迁移集少量样本微调的迁移学习流程。模型先在源域训练集上学习通用特征，再使用迁移集30%的支持样本进行微调，最后在剩余70%的迁移查询样本上评估。")
    add_para(doc, "迁移学习实验显示，迁移查询集Accuracy由0.6461提升至0.6734，Macro F1由0.3122提升至0.3310。正常类、缝内暂堵和缝口暂堵均有改善，说明少量目标段样本能够帮助模型适配新段分布；但其他、延伸受阻等极少样本标签仍未被有效学习，说明迁移学习不能替代基础标签数据补充。")
    add_table(doc, ["指标", "微调前", "微调后", "变化"], [
        ["迁移查询集Accuracy", "0.6461", "0.6734", "+0.0273"],
        ["迁移查询集Macro F1", "0.3122", "0.3310", "+0.0188"],
        ["正常类F1", "0.7703", "0.7962", "提升"],
        ["缝内暂堵F1", "0.5482", "0.6308", "提升明显"],
        ["缝口暂堵F1", "0.6062", "0.6370", "小幅提升"],
    ])
    add_picture(doc, "07_pure_normal_segment_ablation.png", "图3-7 纯正常段剔除消融实验", 5.5)
    doc.add_page_break()

    add_heading(doc, "3.10 当前问题、甲方支持需求与下一步计划")
    add_para(doc, "当前模型已经形成可运行的数据处理、训练评估、图表分析和概率输出流程，但距离工程化稳定应用仍存在三类核心问题：一是异常标签样本总量仍不足，尤其是延伸受阻、滤失过大、缝高延伸等长尾类别；二是不同段之间标签分布差异明显，按段划分后部分类别在训练或测试中缺失，导致模型泛化判断不稳定；三是标签质量和边界一致性仍需专家复核，异常起止点若存在偏差，会直接影响模型学习到的工况边界。")
    add_para(doc, "因此，下一阶段不建议仅依赖调参提升指标，而应围绕数据闭环、标签复核和在线验证推进。模型侧继续保留GNN作为时序图结构主线，同时使用LightGBM等强基线进行结果校验；业务侧需要甲方提供更多异常段、明确标签定义、组织专家复核典型误判样本，并提供跨井、跨平台或新施工批次数据用于独立验证。")
    add_bullets(doc, [
        "数据补充：优先补充每类异常至少覆盖多个井段的样本，避免标签只出现在个别段。",
        "标签复核：对主缝延伸、缝内暂堵、缝口暂堵、砂堵等容易混淆类别建立统一判定口径。",
        "在线验证：将模型输出从硬分类扩展为概率和Top-K风险提示，由现场专家确认是否可用。",
        "模型迭代：继续开展GNN时序图模型、转移概率模型、迁移学习和小样本类重采样策略对比。",
        "验收指标：建议按异常召回、异常Precision、Macro F1、分段泛化和提前预警有效性综合验收。",
    ])
    add_table(doc, ["下一步任务", "需要甲方支持", "预期产出"], [
        ["长尾异常补样", "提供更多含异常的原始段数据和专家确认标签", "提升多分类异常识别稳定性"],
        ["误判样本复核", "安排专家确认模型误报/漏报片段", "修正标签边界和类别口径"],
        ["迁移验证", "提供独立井段或新施工批次数据", "验证跨段、跨井泛化能力"],
        ["预警规则联调", "明确可接受误报率和提前量", "形成现场可解释风险提示"],
    ])

    doc.save(OUT_DOCX)


MD_TEXT = """# 3 基于GNN的压裂工况标签识别研究

## 3.1 研究目标与总体思路

本部分围绕压裂施工秒点数据的工况标签识别开展研究，目标是在已有专家规则标注和现场标签基础上，建立能够利用时间连续性、段内结构关系和动态工艺参数变化的智能识别模型。模型输出对象为 `WORKING_TYPE`，即压裂施工过程中的工况标签；输入对象为砂比、泵压、排量、粘度、累计液量等工艺数值特征及其时间变化特征。

前期探索从原始多分类GNN开始，随后逐步增加正常/异常二分类、异常样本重权重、加砂后截断、动态特征、LightGBM强基线、多分类恢复、下一工况转移概率预测以及迁移学习实验，形成从标签识别到风险概率输出的完整实验链路。

## 3.2 数据基础与标签体系

原始数据约16.40万条，异常占比约1.25%；新增便签数据3.38万条，异常占比约23.16%；合并后总样本约19.78万条，异常占比提升至约5.00%。新增数据补充了缝内暂堵、缝高延伸、延伸受阻、滤失过大等工况，使多分类体系更完整。

## 3.3 样本组织与GNN建模方式

每个压裂段是一条按时间排序的10秒采样序列，段与段之间不拼接。训练、验证、测试和迁移学习集均按段划分，避免同段相邻点同时出现在训练和测试中造成时序泄漏。GNN中节点对应10秒采样点，边对应时间相邻关系，节点特征为工艺数值及动态特征，标签为目标点 `WORKING_TYPE`。

## 3.4 类别不平衡与小样本问题诊断

原始数据中异常占比极低，模型容易趋向输出正常。部分异常类型只出现在少数段甚至个别段中，按段划分后可能出现训练集或测试集中类别缺失的问题。因此评估不能只看Accuracy，应重点看Macro F1、异常类Precision、异常类Recall和混淆矩阵。

## 3.5 数据处理与特征工程改进

实现了加砂后截断、正常下采样、类别权重和动态特征。动态特征包括一阶差分、二阶差分、滚动均值、滚动标准差、窗口最大最小值、斜率、砂比正值比例等，用于表达参数变化方向和变化速度。

## 3.6 模型实验结果与阶段进展

原始GNN在不平衡数据上异常识别较弱；二分类GNN测试Macro F1约0.517；LightGBM二分类约0.524；加砂后截断与动态特征后，二分类最优测试Macro F1提升至约0.5533；合并新增数据并恢复多分类后，窗口4测试Macro F1达到0.5739。

## 3.7 多分类识别效果分析

合并数据多分类实验中，窗口4测试集Accuracy为0.8070，Macro F1为0.5739。正常类F1为0.8942，缝口暂堵F1为0.6097，缝内暂堵F1为0.4978，主缝延伸F1为0.2940。模型已经不再是简单“全预测正常”，但长尾类别仍需要更多数据支撑。

## 3.8 下一工况转移概率预测

建立严格的“当前工况到下一工况”预测模型。输入为当前时刻及历史窗口工艺特征，并显式输入当前工况 `y_t`，输出下一采样点工况 `y_{t+1}` 的概率分布。测试集Accuracy为0.9947，Macro F1为0.9846，迁移集Macro F1为0.8106。该模型适合解释为短时转移风险模型，而不是纯标签识别模型。

## 3.9 迁移学习与跨段泛化

迁移学习采用源域预训练和迁移集少量支持样本微调。迁移查询集Accuracy由0.6461提升至0.6734，Macro F1由0.3122提升至0.3310。结果说明少量目标段样本能够帮助模型适配新段分布，但不能替代长尾异常数据补充。

## 3.10 当前问题、甲方支持需求与下一步计划

当前主要问题包括异常样本不足、长尾类别分布集中、跨段泛化不稳定、标签边界需要专家复核。下一步应优先补充多井段异常样本，统一标签判定口径，开展误判样本专家复核，使用独立井段验证迁移能力，并将模型输出从硬分类扩展为概率和Top-K风险提示。
"""


def build_md() -> None:
    OUT_MD.write_text(MD_TEXT, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(OUT_DOCX)
    print(OUT_MD)
