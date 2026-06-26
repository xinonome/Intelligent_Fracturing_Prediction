from __future__ import annotations

from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
FIG_DIR = ROOT / "report_figures"
OUT_DOCX = ROOT / "第三部分_基于GNN的压裂工况标签识别阶段总结_扩展版.docx"
OUT_MD = ROOT / "第三部分_基于GNN的压裂工况标签识别阶段总结_扩展版.md"


def setup_font() -> None:
    plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "Arial Unicode MS", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False


def save_extra_figures() -> None:
    FIG_DIR.mkdir(exist_ok=True)
    setup_font()

    labels = ["原始数据", "新增便签数据", "合并数据"]
    normal = np.array([161900, 25972, 187872])
    abnormal = np.array([2055, 7828, 9883])
    x = np.arange(len(labels))
    fig, ax = plt.subplots(figsize=(8, 4.6), dpi=160)
    ax.bar(x, normal, label="正常", color="#5B8FF9")
    ax.bar(x, abnormal, bottom=normal, label="异常", color="#F4664A")
    ax.set_xticks(x, labels)
    ax.set_ylabel("样本数")
    ax.set_title("数据补充前后正常/异常样本规模变化")
    ax.legend()
    for i, (n, a) in enumerate(zip(normal, abnormal)):
        ax.text(i, n + a + 2500, f"异常占比 {a/(n+a):.2%}", ha="center", fontsize=9)
    fig.tight_layout()
    fig.savefig(FIG_DIR / "08_dataset_growth_stacked.png", bbox_inches="tight")
    plt.close(fig)

    steps = ["分段数据", "字段筛选", "窗口构样", "GNN/基线训练", "多分类识别", "转移概率", "迁移学习"]
    fig, ax = plt.subplots(figsize=(9, 2.2), dpi=160)
    ax.axis("off")
    xs = np.linspace(0.06, 0.94, len(steps))
    for i, (pos, step) in enumerate(zip(xs, steps)):
        ax.text(pos, 0.5, step, ha="center", va="center", fontsize=10,
                bbox=dict(boxstyle="round,pad=0.35", fc="#E8F3FF", ec="#3A6EA5"))
        if i < len(steps) - 1:
            ax.annotate("", xy=(xs[i + 1] - 0.055, 0.5), xytext=(pos + 0.055, 0.5),
                        arrowprops=dict(arrowstyle="->", lw=1.4, color="#3A6EA5"))
    ax.set_title("压裂工况标签识别建模流程", fontsize=13, pad=12)
    fig.tight_layout()
    fig.savefig(FIG_DIR / "09_modeling_pipeline.png", bbox_inches="tight")
    plt.close(fig)

    metrics = ["Accuracy", "Macro F1"]
    before = [0.6461, 0.3122]
    after = [0.6734, 0.3310]
    x = np.arange(len(metrics))
    width = 0.34
    fig, ax = plt.subplots(figsize=(6.5, 4.2), dpi=160)
    ax.bar(x - width / 2, before, width, label="微调前", color="#7A869A")
    ax.bar(x + width / 2, after, width, label="微调后", color="#36B37E")
    ax.set_xticks(x, metrics)
    ax.set_ylim(0, 0.8)
    ax.set_ylabel("指标值")
    ax.set_title("迁移学习微调前后效果对比")
    ax.legend()
    for i, (b, a) in enumerate(zip(before, after)):
        ax.text(i - width / 2, b + 0.025, f"{b:.4f}", ha="center", fontsize=9)
        ax.text(i + width / 2, a + 0.025, f"{a:.4f}", ha="center", fontsize=9)
    fig.tight_layout()
    fig.savefig(FIG_DIR / "10_transfer_learning_gain.png", bbox_inches="tight")
    plt.close(fig)

    fig, ax = plt.subplots(figsize=(8.2, 3.0), dpi=160)
    ax.axis("off")
    boxes = [
        ("历史窗口特征\nX[t-3:t]", 0.13),
        ("当前工况\ny[t]", 0.38),
        ("转移概率模型", 0.62),
        ("下一工况概率\nP(y[t+1]|X,y[t])", 0.86),
    ]
    for text, xpos in boxes:
        fc = "#FFF4E6" if "当前工况" in text else "#EAF7EA" if "概率" in text else "#E8F3FF"
        ax.text(xpos, 0.5, text, ha="center", va="center", fontsize=11,
                bbox=dict(boxstyle="round,pad=0.45", fc=fc, ec="#2F5597"))
    for a, b in [(0.23, 0.32), (0.47, 0.54), (0.70, 0.78)]:
        ax.annotate("", xy=(b, 0.5), xytext=(a, 0.5),
                    arrowprops=dict(arrowstyle="->", lw=1.5, color="#2F5597"))
    ax.set_title("严格下一工况预测任务定义", fontsize=13, pad=12)
    fig.tight_layout()
    fig.savefig(FIG_DIR / "11_next_state_task.png", bbox_inches="tight")
    plt.close(fig)


def set_fonts(run, name="宋体", size=10.5, bold=False) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_doc_style(doc: Document) -> None:
    for style_name in ["Normal", "Body Text"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "宋体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "黑体"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            style.font.color.rgb = RGBColor(0, 0, 0)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    set_fonts(run, "黑体", 14 if level == 1 else 12, True)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text)
    set_fonts(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        set_fonts(run)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_fonts(run, size=9.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell_text(table.rows[0].cells[i], header, True)
        shade(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value)
    doc.add_paragraph()


def add_picture(doc: Document, filename: str, caption: str, width: float = 5.8) -> None:
    path = FIG_DIR / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_fonts(run, size=9)


def build_docx() -> None:
    save_extra_figures()
    doc = Document()
    set_doc_style(doc)
    section = doc.sections[0]
    section.page_width = 7560310
    section.page_height = 10692130
    section.left_margin = 720090
    section.right_margin = 720090
    section.top_margin = 720090
    section.bottom_margin = 720090

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("3  基于GNN的压裂工况标签识别研究")
    set_fonts(run, "黑体", 18, True)
    doc.add_paragraph()

    add_heading(doc, "3.1 研究背景、目标与总体思路")
    add_para(doc, "压裂施工过程中的工况变化具有明显的时序性和阶段性。现场采集的秒点数据能够连续反映泵压、排量、砂比、粘度、累计砂量、累计液量等参数变化，但这些参数与复杂工况标签之间并不是简单的一一对应关系。某些异常工况表现为趋势变化而非单点突变，例如排量、压力和砂比在一段时间内共同变化后才形成可识别特征；某些工况又具有强烈的段内连续性，相邻采样点之间标签通常不会频繁跳变。因此，基于单点阈值或静态分类的方式难以完整描述施工状态演化，需要引入能够利用时间邻接关系和历史窗口信息的图结构模型。")
    add_para(doc, "本部分工作的核心目标是建立面向压裂秒点数据的工况标签识别模型。模型输出字段固定为WORKING_TYPE，即当前或目标时刻的工况标签；模型输入采用工艺数值特征及其动态变化特征，不使用井名、段名、时间戳、备注、标记字段和结果性字段作为模型输入，以降低信息泄漏风险。建模过程中将每个压裂段视为一组连续数据，段内按照10秒采样点构造图节点，节点之间按照时间相邻关系连边，从而将“压裂段”转化为具有时序结构的图。")
    add_para(doc, "项目推进并不是一次性训练单个模型，而是围绕数据问题逐步迭代。首先用原始多分类GNN验证可行性，随后发现模型受正常类主导，容易产生“几乎全预测正常”的倾向；然后将任务降为正常/异常二分类，并尝试异常过采样、类别权重和正常样本下采样；进一步根据现场业务逻辑剔除加砂前大量正常背景样本，并增加动态特征；在此基础上引入LightGBM作为强基线，与GNN结果交叉验证；最后恢复多分类识别，增加严格的下一工况转移概率预测和迁移学习实验。")
    add_table(doc, ["研究模块", "解决的问题", "当前状态"], [
        ["GNN标签识别", "利用段内时间图结构识别WORKING_TYPE", "已完成多窗口实验和结果评估"],
        ["不平衡处理", "缓解正常样本压制异常样本的问题", "已实现二分类、重权重、下采样、加砂后截断"],
        ["动态特征工程", "表达压力、排量、砂比等参数变化趋势", "已加入差分、滚动统计、斜率等特征"],
        ["多分类恢复", "识别具体异常类型而非仅判断是否异常", "已在合并数据上完成窗口4实验"],
        ["转移概率预测", "输出从当前工况到下一工况的概率", "已完成严格下一状态预测模型"],
        ["迁移学习", "提升新段、新井或新批次数据适配能力", "已完成源域预训练和目标域微调实验"],
    ])
    add_picture(doc, "09_modeling_pipeline.png", "图3-1 压裂工况标签识别建模流程", 5.9)
    doc.add_page_break()

    add_heading(doc, "3.2 数据来源、字段口径与标签体系")
    add_para(doc, "当前数据由两部分组成：一部分是原始分段压裂数据，另一部分是后续加入的便签数据。原始数据包含约16.40万条10秒采样记录，其中正常样本约16.19万条，异常样本约2055条，异常占比仅约1.25%。该比例说明原始训练任务具有极强的类别不平衡特征，如果直接训练多分类模型，损失函数会被正常类主导，模型即使大量预测正常，也能获得较高的表观准确率。")
    add_para(doc, "新增便签数据共约3.38万条记录，其中异常样本约7828条，异常占比约23.16%。这批数据的价值不只是增加了样本量，更重要的是补充了原始数据中缺失或极少出现的工况类型，例如缝内暂堵、缝高延伸、延伸受阻、滤失过大等。合并后总样本量约19.78万条，异常样本约9883条，异常占比提升至约5.00%，数据结构相较原始数据更适合开展多分类识别。")
    add_para(doc, "字段口径方面，WORKING_TYPE是唯一目标标签。空白工况统一视为“正常”，这是后续所有统计、训练和评估的前提。输入特征限定为工艺数值列及其派生特征，包括砂比、泵压、排量、粘度、累计砂量、累计液量、暂堵剂相关参数等。段号、井号、时间、备注、人工标记、结果性字段不作为输入，避免模型学习到与真实工况机理无关的标识信息。")
    add_table(doc, ["数据范围", "总样本量", "正常样本", "异常样本", "异常占比", "主要作用"], [
        ["原始分段数据", "163955", "161900", "2055", "1.25%", "提供基础施工过程和正常背景"],
        ["新增便签数据", "33800", "25972", "7828", "23.16%", "补充异常类型和异常片段"],
        ["合并数据", "197755", "187872", "9883", "5.00%", "支撑多分类和迁移学习实验"],
    ])
    add_picture(doc, "08_dataset_growth_stacked.png", "图3-2 数据补充前后正常/异常样本规模变化", 5.7)
    add_picture(doc, "02_abnormal_label_distribution.png", "图3-3 合并数据异常标签数量分布", 5.5)
    doc.add_page_break()

    add_heading(doc, "3.3 按段划分与窗口样本构造")
    add_para(doc, "压裂施工数据不能简单打乱后随机划分。原因是同一压裂段内相邻采样点高度相关，如果随机划分，训练集可能包含某段的前后点，而测试集包含同一段中间点，模型会在评估阶段看到与训练样本几乎相同的上下文，导致测试指标虚高。为避免这种时序泄漏，本项目坚持按段进行训练集、验证集、测试集和迁移学习集划分。")
    add_para(doc, "样本构造采用滑动窗口方式。对于窗口长度k，模型输入为同一段内连续k个采样点的特征序列，输出为目标点的WORKING_TYPE。窗口长度从1到6分别对应10秒到60秒的历史上下文。所有窗口都必须来自同一压裂段，禁止跨段拼接；若某段长度小于窗口长度，则该段在该窗口设置下不生成样本。")
    add_para(doc, "窗口设计的意义在于比较不同时间上下文对工况识别的贡献。1点输入只使用当前瞬时参数，适合识别参数差异明显的工况；2至3点输入开始引入短时变化；4点输入约覆盖40秒历史，在当前实验中表现最好；5至6点输入虽然包含更长上下文，但也可能引入早期无关波动，导致泛化能力下降。")
    add_table(doc, ["窗口长度", "时间范围", "输入含义", "当前观察"], [
        ["1", "10秒", "单点工艺状态", "容易受瞬时噪声影响"],
        ["2", "20秒", "短时变化", "二分类结果有所改善"],
        ["3", "30秒", "短趋势窗口", "部分实验不稳定"],
        ["4", "40秒", "趋势与噪声平衡", "当前多分类最优窗口"],
        ["5", "50秒", "较长上下文", "可能引入无关片段"],
        ["6", "60秒", "最长窗口", "不一定优于窗口4"],
    ])
    doc.add_page_break()

    add_heading(doc, "3.4 GNN建模方法与算法解释")
    add_para(doc, "GNN建模的基本思想是将压裂段视为时间图。图中的每个节点对应一个10秒采样点，节点特征为该采样点的工艺参数向量；相邻采样点之间建立边，表示施工状态随时间连续演化。对于窗口样本，可以将窗口内节点组成一个局部子图，模型通过图卷积或消息传递聚合相邻节点信息，再对目标节点或窗口整体进行分类。")
    add_para(doc, "这种建模方式相比普通表格分类有两个优势。第一，它保留了段内顺序关系，使模型能够学习“当前点与前序点之间的变化”而不只是当前数值；第二，它天然适合扩展到下一工况转移预测，因为当前节点状态、历史邻居状态和下一节点标签之间可以统一表示为图上的状态传播问题。")
    add_para(doc, "在实验中，GNN并不是唯一模型。我们同时保留LightGBM作为强基线，原因是LightGBM对表格特征、非线性分裂和小样本不平衡问题具有较强鲁棒性。如果GNN显著优于LightGBM，说明图结构确实带来收益；如果LightGBM表现更稳定，则说明当前主要瓶颈可能在数据质量、标签分布或特征口径，而不是模型结构本身。这个对照有助于向算法评审解释：当前工作不是单纯追求复杂模型，而是在用可解释基线校验GNN路线的有效性。")
    add_table(doc, ["模型", "输入形式", "优势", "局限"], [
        ["GNN", "段内时间图/窗口子图", "保留时序邻接关系，适合状态演化", "对样本量和标签质量敏感"],
        ["LightGBM", "窗口展开后的表格特征", "小样本和非线性特征表现稳定", "不能显式表达图结构"],
        ["转移概率模型", "历史窗口+当前工况", "输出下一工况概率，适合预警", "高指标受状态连续性影响"],
        ["迁移学习模型", "源域预训练+目标域微调", "适合新段适配", "需要少量目标域标签"],
    ])
    doc.add_page_break()

    add_heading(doc, "3.5 类别不平衡、小样本与长尾问题")
    add_para(doc, "最初模型效果不理想的根本原因在于样本分布，而不是训练轮数不够。原始数据中正常样本占比接近98.75%，异常仅占1.25%。在这种情况下，模型只要倾向输出正常，就可以在Accuracy上看起来不差，但异常召回会很低。对于压裂复杂情况识别，漏检异常比少量误报更不可接受，因此必须用Macro F1、异常类Recall和异常类Precision来评价模型，而不能只看总准确率。")
    add_para(doc, "长尾问题还体现在异常标签的段级分布上。某些异常只出现在个别段，按段划分后可能导致训练集中没有该类充分样本，测试集中却出现该类；或者训练集中有该类，但测试集中没有该类，指标无法反映泛化能力。这是现场数据建模常见问题：严格按段划分更接近真实部署，但也会暴露数据覆盖不足。")
    add_para(doc, "因此，我们尝试了三类处理策略：第一，任务降维，将多分类先转为正常/异常二分类，验证模型能否识别异常边界；第二，样本再平衡，包括异常过采样、正常下采样和类别权重；第三，业务截断，去除加砂前大量正常背景点，使训练样本更集中于可能发生复杂工况的有效施工区间。这些策略能够改善模型，但不能替代真实异常案例补充。")
    add_bullets(doc, [
        "Accuracy高不代表模型可用，特别是在正常样本极多的场景下。",
        "Macro F1更能反映多类别整体识别能力，因为每个类别权重相同。",
        "异常Recall代表异常被找回的比例，过低意味着漏检风险高。",
        "异常Precision代表模型报出的异常有多少是真的，过低意味着现场误报压力大。",
        "长尾类别如果只有几十个样本，很难仅靠算法稳定识别。",
    ])
    add_picture(doc, "03_new_segment_label_heatmap.png", "图3-4 各段标签分布差异", 5.8)
    doc.add_page_break()

    add_heading(doc, "3.6 加砂后截断与动态特征工程")
    add_para(doc, "从业务机理看，加砂前大部分时间处于准备或正常施工背景，复杂工况出现概率较低。将这些样本全部纳入训练，会进一步放大正常类比例，使模型更加偏向正常。基于这一判断，我们实现了加砂后截断，即以砂比SB首次大于0作为有效识别区间起点，只保留加砂后的连续施工片段用于异常识别实验。")
    add_para(doc, "需要说明的是，截断不会直接改变原始异常点本身，但会改变窗口样本数量。因为滑动窗口需要连续k个点，当截断起点变化、窗口长度变化、段级划分变化时，可生成的样本数会变化。因此在不同实验中看到异常窗口数变化，是窗口构样逻辑和数据划分共同作用的结果，并不意味着原始异常标签被随意修改。")
    add_para(doc, "动态特征工程是本阶段提升效果最明显的技术措施之一。现场专家识别复杂工况时，通常不会只看某一秒的压力或砂比，而是看一段时间内压力是否持续升高、排量是否下降、砂比是否变化、累计液量和累计砂量是否出现异常组合。因此，我们对核心工艺参数构造差分、滚动均值、滚动标准差、滚动最大最小值、斜率和正值比例等特征，使模型能够学习“变化过程”而不是孤立点值。")
    add_table(doc, ["特征类别", "示例", "业务含义"], [
        ["原始特征", "泵压、排量、砂比、粘度、累计液量", "描述当前施工状态"],
        ["差分特征", "一阶差分、二阶差分", "描述短时上升或下降趋势"],
        ["滚动统计", "3/5/10窗口均值、标准差、最大最小值", "描述局部稳定性和波动"],
        ["斜率特征", "窗口内参数变化斜率", "描述连续变化速度"],
        ["砂比状态", "是否加砂、加砂持续点数、砂比正值比例", "定位有效施工区间"],
    ])
    doc.add_page_break()

    add_heading(doc, "3.7 实验路径与阶段结果")
    add_para(doc, "本阶段形成了多条实验路径。第一条是原始GNN多分类路径，用于验证图神经网络在压裂段时间图上的基本可行性。该路径暴露出明显问题：在极不平衡数据上，模型容易学习到正常类先验，异常类识别能力不足。第二条是正常/异常二分类路径，用于降低任务难度，优先判断模型是否能够识别异常边界。第三条是加砂后截断与动态特征路径，重点改善样本口径和特征表达。第四条是合并新增便签数据后恢复多分类，用于评估具体异常类别识别能力。")
    add_para(doc, "实验结果显示，单纯增加训练轮数不是主要提升来源。原始GNN二分类测试Macro F1约0.517，LightGBM二分类约0.524；加入加砂后截断和动态特征后，二分类最优Macro F1提升至约0.5533；合并新增数据并恢复多分类后，窗口4测试Macro F1达到0.5739。这个结果说明，模型效果提升主要来自三点：异常数据补充、有效施工区间筛选、动态特征表达。")
    add_para(doc, "纯正常段剔除实验也提供了重要结论。合并数据中共有51个段，其中纯正常段8个，含异常段43个。直接剔除纯正常段后，多分类窗口4测试Macro F1从0.5739下降到0.4389，说明纯正常段并非完全无用。它们虽然不包含异常，但为模型提供了正常边界和正常波动范围，有助于减少误报。因此，后续不应简单删除纯正常段，而应通过权重、采样比例或分阶段训练控制其影响。")
    add_table(doc, ["实验阶段", "关键设置", "测试Macro F1", "主要结论"], [
        ["原始GNN二分类", "全量数据，正常/异常", "约0.517", "可识别部分异常，但受不平衡影响"],
        ["LightGBM二分类", "同口径强基线", "约0.524", "表格强基线略稳定"],
        ["加砂后+动态特征", "截断、动态特征、下采样", "约0.5533", "样本口径和特征工程有效"],
        ["合并数据多分类", "新增便签数据，窗口4", "0.5739", "异常数据补充带来提升"],
        ["剔除纯正常段", "仅保留含异常段", "0.4389", "简单删除纯正常段会损害正常边界学习"],
    ])
    add_picture(doc, "04_model_macro_f1_comparison.png", "图3-5 不同实验路径Macro F1对比", 5.8)
    add_picture(doc, "07_pure_normal_segment_ablation.png", "图3-6 纯正常段剔除消融实验", 5.5)
    doc.add_page_break()

    add_heading(doc, "3.8 多分类识别效果与指标解释")
    add_para(doc, "合并数据多分类实验中，当前最优窗口为4，测试集Accuracy为0.8070，Macro F1为0.5739。正常类F1为0.8942，说明模型对主类识别较稳定；缝口暂堵F1为0.6097，缝内暂堵F1为0.4978，说明新增数据中的暂堵类已经具备一定可学习性；主缝延伸Precision达到0.9365，但Recall仅0.1744，说明模型一旦预测主缝延伸通常较准，但漏检较多。")
    add_para(doc, "这里需要向汇报对象明确Precision、Recall和F1的含义。Precision表示模型报出某一类时有多少是真的，适合衡量误报；Recall表示真实属于某一类的样本有多少被模型找回，适合衡量漏报；F1是Precision和Recall的调和平均，计算公式为F1=2×Precision×Recall/(Precision+Recall)。例如某类Precision高但Recall低，说明模型判断谨慎，误报少但漏报多；某类Recall高但Precision低，则说明模型报得多，漏检少但误报多。")
    add_para(doc, "当前多分类结果说明模型已经不再是“几乎全预测正常”。它能够识别缝口暂堵、缝内暂堵等部分类别，但对主缝延伸存在漏检，对其他、延伸受阻、滤失过大、缝高延伸等极少样本类别还不能形成稳定结论。后续若要把模型从阶段性研究推进到工程应用，必须将评价从整体Macro F1进一步拆解到每个异常类型的Recall、Precision和可接受误报率。")
    add_table(doc, ["类别", "Precision", "Recall", "F1", "测试样本数", "解释"], [
        ["正常", "0.8321", "0.9663", "0.8942", "4570", "识别稳定，但仍有异常被吸收到正常类"],
        ["主缝延伸", "0.9365", "0.1744", "0.2940", "1015", "预测较准，但漏检明显"],
        ["缝内暂堵", "0.4130", "0.6264", "0.4978", "182", "召回尚可，误报偏多"],
        ["缝口暂堵", "0.6157", "0.6038", "0.6097", "260", "Precision和Recall相对均衡"],
    ])
    add_picture(doc, "05_multiclass_per_class_metrics.png", "图3-7 多分类测试集各类别识别指标", 5.8)
    doc.add_page_break()

    add_heading(doc, "3.9 下一工况转移概率预测")
    add_para(doc, "除当前标签识别外，现场更关注复杂工况是否即将发生。因此，我们进一步建立严格的下一工况预测模型。该模型的目标不是预测当前点标签，而是在已知当前工况y_t和历史窗口特征X[t-k+1:t]的条件下，预测下一采样点y_{t+1}属于各类工况的概率。输出形式不是单一类别，而是一组概率，例如当前为正常时，下一点仍为正常、转为砂堵、转为缝内暂堵、转为滤失过大等概率。")
    add_para(doc, "该任务必须严格区分输入和输出。输入可以包含当前工况，因为业务问题本身就是“从当前工况转移到下一工况”；输出只能是下一时刻标签，不能把未来点的参数或未来标签泄漏给模型。当前实现中，训练样本按段内时间顺序构造，跨段位置不生成转移样本，保证转移概率只描述同一压裂段内相邻采样点的状态演化。")
    add_para(doc, "转移模型测试集Accuracy达到0.9947，Macro F1达到0.9846，迁移集Macro F1达到0.8106。这个结果较高，主要原因是压裂工况具有强连续性，绝大多数相邻10秒点状态不发生变化。因此，该模型不能简单解释为“异常识别已经接近完美”，而应解释为短时状态保持与转移风险估计模型。它的价值在于输出概率排序，用于提示下一时刻从正常转向某类异常的风险，而不是替代当前点识别模型。")
    add_table(doc, ["任务", "输入", "输出", "适用场景"], [
        ["当前标签识别", "历史窗口工艺特征", "当前WORKING_TYPE", "离线标注、当前状态判断"],
        ["下一工况预测", "历史窗口工艺特征+当前工况", "下一点各类工况概率", "短时风险提示、状态转移分析"],
        ["经验转移统计", "历史标签序列", "经验转移频率", "规则校验、模型概率解释"],
    ])
    add_picture(doc, "11_next_state_task.png", "图3-8 严格下一工况预测任务定义", 5.8)
    add_picture(doc, "06_normal_to_abnormal_transition_probs.png", "图3-9 正常状态转向异常状态的预测概率", 5.8)
    doc.add_page_break()

    add_heading(doc, "3.10 迁移学习实验与跨段泛化")
    add_para(doc, "压裂数据存在明显的跨段和跨井差异。不同段的地质条件、施工阶段、参数范围和异常类型不完全一致，因此在已有段上训练出的模型，直接应用到新段时往往会出现性能下降。为了验证少量目标段标签对模型适配的作用，本阶段实现了迁移学习流程：先在源域训练集上进行预训练，再取迁移集30%的支持样本进行微调，最后在迁移集剩余70%的查询样本上评估。")
    add_para(doc, "实验结果显示，迁移查询集Accuracy由0.6461提升到0.6734，Macro F1由0.3122提升到0.3310。正常类、缝内暂堵和缝口暂堵均有改善，说明少量目标域样本确实能够帮助模型适配新段分布。尤其是缝内暂堵F1从0.5482提升到0.6308，说明在已有一定样本基础上，目标段微调可以提升局部异常识别能力。")
    add_para(doc, "但迁移学习也暴露出限制。对于其他、延伸受阻等极少样本类别，微调后仍不能可靠识别；主缝延伸虽然Recall提高，但Precision下降，说明模型在目标域中更积极地报出该类，带来了误报风险。因此，迁移学习适合作为新井段上线前的适配手段，但仍需要基础数据覆盖和专家标签质量作为前提。")
    add_table(doc, ["指标", "微调前", "微调后", "变化", "解释"], [
        ["迁移查询集Accuracy", "0.6461", "0.6734", "+0.0273", "整体预测准确率提升"],
        ["迁移查询集Macro F1", "0.3122", "0.3310", "+0.0188", "类别均衡指标小幅提升"],
        ["正常类F1", "0.7703", "0.7962", "提升", "目标域正常边界更清晰"],
        ["缝内暂堵F1", "0.5482", "0.6308", "提升明显", "目标域异常适配有效"],
        ["缝口暂堵F1", "0.6062", "0.6370", "小幅提升", "已有类别进一步稳定"],
    ])
    add_picture(doc, "10_transfer_learning_gain.png", "图3-10 迁移学习微调前后效果对比", 5.2)
    doc.add_page_break()

    add_heading(doc, "3.11 当前结论、问题边界与甲方支持需求")
    add_para(doc, "从目前结果看，项目已经完成了从数据清洗、字段口径确认、分段划分、窗口构样、GNN训练、强基线对照、多分类恢复、转移概率预测到迁移学习验证的完整闭环。阶段性结论可以概括为：模型已具备识别部分异常工况的能力，新增异常数据和动态特征是提升效果的主要来源；但当前精度仍不足以直接作为现场自动决策依据，更适合定位为辅助标注、风险提示和专家复核工具。")
    add_para(doc, "当前最大瓶颈是数据而不是单纯算法。异常样本总量不足、长尾类别只出现在个别段、部分类别训练集和测试集分布不一致、标签起止边界缺少统一复核，都会限制模型上限。对于算法评审而言，需要明确：继续调参可以带来有限提升，但如果没有更多覆盖不同井段、不同异常类型、不同施工阶段的高质量标注数据，模型很难稳定识别长尾工况。")
    add_para(doc, "向甲方需要争取的支持包括三方面。第一，数据支持：补充更多含异常的原始段数据，尤其是延伸受阻、滤失过大、缝高延伸、砂堵等低频工况，并尽量保证每类异常覆盖多个井段。第二，专家支持：对典型误判、漏判片段进行复核，明确各类工况的判定边界和标签起止点。第三，验证支持：提供独立井段或新施工批次数据，用于验证模型跨段、跨井和跨批次泛化能力。")
    add_table(doc, ["问题", "对模型的影响", "需要支持"], [
        ["异常样本少", "长尾类别难以学习", "补充更多异常段和异常标签"],
        ["标签边界不一致", "模型学习到模糊边界", "专家复核起止点"],
        ["跨段分布差异", "迁移集效果下降", "提供独立井段验证集"],
        ["误报/漏报权衡不明确", "阈值无法业务化", "明确现场可接受误报率和提前量"],
    ])
    doc.add_page_break()

    add_heading(doc, "3.12 下一步工作计划")
    add_para(doc, "下一阶段建议按照“数据闭环优先、模型迭代跟进、现场验证落地”的路线推进。首先建立异常样本补充和专家复核机制，对每类异常形成最小可训练样本集；其次继续优化GNN结构，引入边权、节点时间编码和状态转移约束，使模型更好地利用段内图结构；再次将当前识别模型与下一工况转移概率模型结合，形成当前状态识别、下一状态概率、Top-K风险提示三类输出。")
    add_para(doc, "在工程化方向上，应逐步从离线批处理走向准实时应用。模型每10秒接收最新采样点，更新历史窗口特征和图节点状态，输出当前工况标签及下一工况概率。当异常概率超过阈值时，不直接替代人工判断，而是向现场提供风险提示、可能异常类别、关键特征变化和历史相似片段，辅助专家快速确认。")
    add_para(doc, "验收指标也应从单一准确率改为综合指标。建议同时考察正常/异常二分类Recall、多分类Macro F1、重点异常类Recall、误报率、提前预警命中率、跨段泛化能力和专家复核通过率。只有当模型在独立井段上保持稳定表现，并且误报与漏报达到现场可接受范围后，才适合进入更深层次的自动化应用。")
    add_bullets(doc, [
        "短期：补充异常样本，完善标签口径，固定训练/验证/测试/迁移划分规则。",
        "中期：优化GNN结构，融合转移概率输出，建立误判样本回流机制。",
        "长期：接入实时数据流，实现在线风险提示和专家反馈闭环。",
        "汇报建议：当前应强调模型已证明路线可行，但精度瓶颈主要来自异常样本覆盖和标签质量，需要甲方数据与专家资源支持。",
    ])

    doc.save(OUT_DOCX)


MD_TEXT = """# 3 基于GNN的压裂工况标签识别研究

本扩展版报告围绕研究目标、数据基础、GNN建模、类别不平衡处理、动态特征、多分类识别、下一工况转移概率、迁移学习和下一步计划展开。Word版本中已包含完整正文、表格和配图。

核心结论：

- 原始数据异常占比约1.25%，模型容易偏向正常类。
- 新增便签数据后，合并数据异常占比提升至约5.00%，补充了缝内暂堵、缝高延伸、延伸受阻、滤失过大等工况。
- 加砂后截断和动态特征使二分类最优Macro F1提升至约0.5533。
- 合并数据多分类窗口4测试Macro F1达到0.5739，正常、缝口暂堵、缝内暂堵已有一定识别能力，但主缝延伸漏检较多，长尾类别仍不足。
- 严格下一工况转移概率模型测试Macro F1达到0.9846，但该结果主要反映状态连续性，应解释为短时转移风险模型。
- 迁移学习微调后，迁移查询集Accuracy由0.6461提升到0.6734，Macro F1由0.3122提升到0.3310。
- 后续提升的关键不是单纯调参，而是补充异常样本、统一标签边界、提供独立井段验证集和专家复核机制。
"""


def build_md() -> None:
    OUT_MD.write_text(MD_TEXT, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_md()
    print(OUT_DOCX)
    print(OUT_MD)
